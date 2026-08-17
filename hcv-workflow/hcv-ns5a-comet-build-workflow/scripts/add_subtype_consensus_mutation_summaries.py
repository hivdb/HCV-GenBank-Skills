#!/usr/bin/env python3
"""Add reference-to-consensus mutation summaries to subtype comparison workbooks."""
from __future__ import annotations
import argparse
import re
import shutil
from pathlib import Path
from docx import Document
from openpyxl import load_workbook
from openpyxl.cell.rich_text import CellRichText, TextBlock
from openpyxl.cell.text import InlineFont
from openpyxl.styles import Font

VARIANT = re.compile(r"([A-Z*])(\d+(?:\.\d+)?)")
PERCENTAGE = re.compile(r"(?:\d+(?:\.\d+)?|NA)%")
MUTATION_POSITION = re.compile(r"^[A-Z*](\d+)")

def profiles(path: Path) -> dict[tuple[str,str,int,str],str]:
    wb=load_workbook(path,read_only=True,data_only=True); ws=wb.active; result={}
    header=[c.value for c in next(ws.iter_rows(min_row=1,max_row=1))]
    for row in ws.iter_rows(min_row=2,values_only=True):
        label=str(row[0] or ''); m=re.match(r"GT(\d+)_(\S+) \(",label)
        if not m: continue
        for column,value in enumerate(row[1:],1):
            if column>=len(header) or not str(header[column] or '').startswith('P'): continue
            pos=int(str(header[column])[1:])
            for aa,pct in VARIANT.findall(str(value or '')): result[(m.group(1),m.group(2).lower(),pos,aa)]=pct
    wb.close(); return result


def mutation_sort_key(value: str) -> tuple[int, str]:
    match = MUTATION_POSITION.match(value)
    return (int(match.group(1)) if match else 1_000_000, value)

def write_summary_documents(root: Path, readme: dict[str,dict[str,set[str]]]) -> None:
    title = 'Subtype reference-to-COMET-consensus mutations'
    note = 'Only subtypes with one or more RAS-position differences are listed. Percentages are from the final combined profile.'
    lines=[f'# {title}', '', note, '']
    document = Document()
    document.add_heading(title, level=0)
    document.add_paragraph(note)
    for gene in ('NS3','NS5A_NTD','NS5B'):
        lines.extend([f'## {gene}', ''])
        document.add_heading(gene, level=1)
        for subtype, mutations in sorted(readme.get(gene,{}).items(), key=lambda item: (int(item[0][2]), item[0])):
            summary = f'{subtype}: {", ".join(sorted(mutations, key=mutation_sort_key))}'
            lines.append(f'- {summary}')
            paragraph = document.add_paragraph(style='List Bullet')
            offset = 0
            for match in PERCENTAGE.finditer(summary):
                paragraph.add_run(summary[offset:match.start()])
                percentage_run = paragraph.add_run(match.group())
                percentage_run.font.subscript = True
                offset = match.end()
            paragraph.add_run(summary[offset:])
        lines.append('')
    (root/'README_Subtype_Consensus_Mutations.md').write_text('\n'.join(lines),encoding='utf-8')
    document.save(root/'README_Subtype_Consensus_Mutations.docx')

def publish_shared_report(source: Path, destination: Path) -> None:
    """Copy the complete report set after all three COMET workflows are available."""
    destination.mkdir(parents=True, exist_ok=True)
    filenames = [
        'HCV_Subtype_Ref_vs_Comet_Subtype_Consensus_Aligned_NS3.xlsx',
        'HCV_Subtype_Ref_vs_Comet_Subtype_Consensus_Aligned_NS5A_NTD.xlsx',
        'HCV_Subtype_Ref_vs_Comet_Subtype_Consensus_Aligned_NS5B.xlsx',
        'README_Subtype_Consensus_Mutations.md',
        'README_Subtype_Consensus_Mutations.docx',
    ]
    for filename in filenames:
        shutil.copy2(source / filename, destination / filename)


def validate_reported_mutations(
    output_root: Path,
    comet: Path,
    gene_tokens: list[tuple[str, str]],
) -> None:
    """Require every reported RAS consensus mutation to have a profile percentage."""
    missing: list[str] = []
    for gene, token in gene_tokens:
        values = profiles(comet / f"{token}_Subtype_RAS_Profiles.xlsx")
        workbook = load_workbook(
            output_root / f"HCV_Subtype_Ref_vs_Comet_Subtype_Consensus_Aligned_{gene}.xlsx",
            read_only=True,
            data_only=True,
        )
        worksheet = workbook.active
        mutation_column = next((cell.column for cell in worksheet[2] if cell.value == "Mutations"), worksheet.max_column + 1)
        for row in range(1, worksheet.max_row + 1):
            if worksheet.cell(row, 4).value != "Consensus":
                continue
            genotype = str(worksheet.cell(row, 2).value or "").removeprefix("GT")
            subtype = str(worksheet.cell(row, 3).value or "").lower()
            for column in range(5, mutation_column):
                consensus_aa = worksheet.cell(row, column).value
                reference_aa = worksheet.cell(row - 1, column).value
                position = worksheet.cell(2, column).value
                if not consensus_aa or not reference_aa or not position:
                    continue
                key = (genotype, subtype, int(position), str(consensus_aa))
                if key not in values:
                    missing.append(f"{gene}:GT{genotype}_{subtype}:{reference_aa}{position}{consensus_aa}")
        workbook.close()
    if missing:
        examples = "; ".join(missing[:20])
        suffix = "" if len(missing) <= 20 else f"; ... ({len(missing)} total)"
        raise RuntimeError(
            "Reported RAS consensus mutations are missing from the subtype RAS profile: "
            f"{examples}{suffix}"
        )

def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        '--shared-report-dir', type=Path,
        default=Path('outputs/shared_report/ICTV_ref_local_cons_compare'),
        help='Destination for copies of the final three-gene report set.',
    )
    args = parser.parse_args()
    output_root=Path('outputs/reference_seqs'); comet=Path('outputs/comet')
    required = [
        *(output_root / f'HCV_Subtype_Ref_vs_Comet_Subtype_Consensus_Aligned_{gene}.xlsx'
          for gene in ('NS3', 'NS5A_NTD', 'NS5B')),
        *(comet / f'{gene}_Subtype_RAS_Profiles.xlsx' for gene in ('NS3', 'NS5A', 'NS5B')),
    ]
    missing = [str(path) for path in required if not path.is_file()]
    if missing:
        print('shared_report_status=waiting_for_all_comet_workflows')
        print('missing_inputs=' + ';'.join(missing))
        return
    gene_tokens = [('NS3','NS3'),('NS5A_NTD','NS5A'),('NS5B','NS5B')]
    validate_reported_mutations(output_root, comet, gene_tokens)
    readme: dict[str,dict[str,set[str]]] = {}
    for gene,token in gene_tokens:
        path=output_root/f'HCV_Subtype_Ref_vs_Comet_Subtype_Consensus_Aligned_{gene}.xlsx'
        values=profiles(comet/f'{token}_Subtype_RAS_Profiles.xlsx')
        wb=load_workbook(path); ws=wb.active
        col=next((cell.column for cell in ws[2] if cell.value=='Mutations'), ws.max_column+1)
        for row in range(1,ws.max_row+1):
            if ws.cell(row,1).value=='Gene': ws.cell(row,col).value='Mutations'; ws.cell(row,col).font=Font(bold=True)
            if ws.cell(row,4).value!='Consensus': continue
            refrow=row-1; gt=str(ws.cell(row,2).value or '').removeprefix('GT'); subtype=str(ws.cell(row,3).value or '').lower(); parts=[]; plain=[]
            for c in range(5,col):
                cons=ws.cell(row,c).value
                if not cons: continue
                ref=ws.cell(refrow,c).value; pos=ws.cell(2,c).value
                if not ref or not pos: continue
                mutation=f'{ref}{pos}{cons}'; pct=values[(gt,subtype,int(pos),str(cons))]
                parts.extend([mutation, TextBlock(InlineFont(vertAlign='subscript'), pct), ', '])
                plain.append(f'{mutation} {pct}%')
            if parts:
                parts.pop(); ws.cell(row,col).value=CellRichText(*parts)
                readme.setdefault(gene,{}).setdefault(f'GT{gt}_{subtype}',set()).update(plain)
        ws.column_dimensions[ws.cell(1,col).column_letter].width=50; wb.save(path)
    write_summary_documents(output_root, readme)
    publish_shared_report(output_root, args.shared_report_dir)

if __name__=='__main__': main()
