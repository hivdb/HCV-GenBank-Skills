#!/usr/bin/env python3
"""Add reference-to-consensus mutation summaries to subtype comparison workbooks."""
from __future__ import annotations
import argparse
import re
import shutil
from pathlib import Path
from docx import Document
from openpyxl import load_workbook
from openpyxl.cell.rich_text import CellRichText, TextBlock
from openpyxl.cell.text import InlineFont
from openpyxl.styles import Font

VARIANT = re.compile(r"([A-Z*])(\d+(?:\.\d+)?)")
PERCENTAGE = re.compile(r"(?:\d+(?:\.\d+)?|NA)%")
MUTATION_POSITION = re.compile(r"^[A-Z*](\d+)")

def profiles(path: Path) -> dict[tuple[str,str,int,str],str]:
    wb=load_workbook(path,read_only=True,data_only=True); ws=wb.active; result={}
    header=[c.value for c in next(ws.iter_rows(min_row=1,max_row=1))]
    for row in ws.iter_rows(min_row=2,values_only=True):
        label=str(row[0] or ''); m=re.match(r"GT(\d+)_(\S+) \(",label)
        if not m: continue
        for column,value in enumerate(row[1:],1):
            if column>=len(header) or not str(header[column] or '').startswith('P'): continue
            pos=int(str(header[column])[1:])
            for aa,pct in VARIANT.findall(str(value or '')): result[(m.group(1),m.group(2).lower(),pos,aa)]=pct
    wb.close(); return result


def mutation_sort_key(value: str) -> tuple[int, str]:
    match = MUTATION_POSITION.match(value)
    return (int(match.group(1)) if match else 1_000_000, value)


def combined_profile_subtypes(path: Path) -> set[str]:
    """Return GT/subtype labels represented in a combined RAS profile."""
    workbook = load_workbook(path, read_only=True, data_only=True)
    worksheet = workbook.active
    labels: set[str] = set()
    for row in worksheet.iter_rows(min_row=2, values_only=True):
        label = str(row[0] or "")
        match = re.match(r"GT(\d+)_([^\s(]+)\s*\(", label)
        if match:
            labels.add(f"GT{match.group(1)}_{match.group(2).lower()}")
    workbook.close()
    return labels

def write_summary_documents(root: Path, readme: dict[str,dict[str,set[str]]], genes: list[str]) -> None:
    note = 'Only subtypes with one or more RAS-position differences are listed. Percentages are from the final combined profile.'
    for gene in genes:
        title = f'{gene} subtype reference-to-COMET-consensus mutations'
        lines = [f'# {title}', '', note, '']
        document = Document()
        document.add_heading(title, level=0)
        document.add_paragraph(note)
        for subtype, mutations in sorted(readme.get(gene,{}).items(), key=lambda item: (int(item[0][2]), item[0])):
            summary = f'{subtype}: {", ".join(sorted(mutations, key=mutation_sort_key))}'
            lines.append(f'- {summary}')
            paragraph = document.add_paragraph(style='List Bullet')
            offset = 0
            for match in PERCENTAGE.finditer(summary):
                paragraph.add_run(summary[offset:match.start()])
                percentage_run = paragraph.add_run(match.group())
                percentage_run.font.subscript = True
                offset = match.end()
            paragraph.add_run(summary[offset:])
        stem = f'README_Subtype_Consensus_Mutations_{gene}'
        (root / f'{stem}.md').write_text('\n'.join(lines) + '\n', encoding='utf-8')
        document.save(root / f'{stem}.docx')
    for suffix in ('.md', '.docx'):
        (root / f'README_Subtype_Consensus_Mutations{suffix}').unlink(missing_ok=True)

def publish_shared_report(comparison_paths: dict[str, Path], documents_root: Path, destination: Path, genes: list[str]) -> None:
    """Copy the requested gene-specific report set."""
    destination.mkdir(parents=True, exist_ok=True)
    for gene in genes:
        path = comparison_paths[gene]
        shutil.copy2(path, destination / path.name)
        for suffix in ('.md', '.docx'):
            filename = f'README_Subtype_Consensus_Mutations_{gene}{suffix}'
            shutil.copy2(documents_root / filename, destination / filename)
    for suffix in ('.md', '.docx'):
        (destination / f'README_Subtype_Consensus_Mutations{suffix}').unlink(missing_ok=True)


def validate_reported_mutations(
    comparison_paths: dict[str, Path],
    profile_paths: dict[str, Path],
    gene_tokens: list[tuple[str, str]],
) -> None:
    """Require every reported RAS consensus mutation to have a profile percentage."""
    missing: list[str] = []
    for gene, token in gene_tokens:
        values = profiles(profile_paths[token])
        workbook = load_workbook(
            comparison_paths[gene],
            read_only=True,
            data_only=True,
        )
        worksheet = workbook.active
        mutation_column = next((cell.column for cell in worksheet[2] if cell.value == "Mutations"), worksheet.max_column + 1)
        for row in range(1, worksheet.max_row + 1):
            if worksheet.cell(row, 4).value != "Consensus":
                continue
            genotype = str(worksheet.cell(row, 2).value or "").removeprefix("GT")
            subtype = str(worksheet.cell(row, 3).value or "").lower()
            for column in range(5, mutation_column):
                consensus_aa = worksheet.cell(row, column).value
                reference_aa = worksheet.cell(row - 1, column).value
                position = worksheet.cell(2, column).value
                if not consensus_aa or not reference_aa or not position:
                    continue
                key = (genotype, subtype, int(position), str(consensus_aa))
                if key not in values:
                    missing.append(f"{gene}:GT{genotype}_{subtype}:{reference_aa}{position}{consensus_aa}")
        workbook.close()
    if missing:
        examples = "; ".join(missing[:20])
        suffix = "" if len(missing) <= 20 else f"; ... ({len(missing)} total)"
        raise RuntimeError(
            "Reported RAS consensus mutations are missing from the subtype RAS profile: "
            f"{examples}{suffix}"
        )

def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        '--shared-report-dir', type=Path,
        default=Path('outputs/shared_report/ICTV_ref_local_cons_compare'),
        help='Destination for copies of the final three-gene report set.',
    )
    parser.add_argument('--comparison-output-dir', type=Path, default=Path('outputs/reference_seqs'))
    parser.add_argument('--ns3-output-dir', type=Path, default=Path('outputs/comet-NS3'))
    parser.add_argument('--ns5a-output-dir', type=Path, default=Path('outputs/comet-NS5A-all-ras'))
    parser.add_argument('--ns5b-output-dir', type=Path, default=Path('outputs/comet-NS5B'))
    parser.add_argument('--comparison-workbook', type=Path, help='Explicit comparison workbook for the selected gene.')
    parser.add_argument('--ras-workbook', type=Path, help='Explicit subtype RAS workbook for the selected gene.')
    parser.add_argument('--combined-profile-workbook', type=Path, help='Optional combined RAS profile limiting reported subtypes.')
    parser.add_argument('--gene', required=True, choices=('NS3', 'NS5A', 'NS5B'), help='COMET workflow gene to publish.')
    args = parser.parse_args()
    output_root = args.comparison_output_dir
    comparison_paths = {
        'NS3': output_root / 'HCV_Subtype_Ref_vs_Comet_Subtype_Consensus_Aligned_NS3.xlsx',
        'NS5A_NTD': output_root / 'HCV_Subtype_Ref_vs_Comet_Subtype_Consensus_Aligned_NS5A_NTD.xlsx',
        'NS5B': args.ns5b_output_dir / 'HCV_Subtype_Ref_vs_Comet_Subtype_Consensus_Aligned_NS5B.xlsx',
    }
    profile_paths = {
        'NS3': args.ns3_output_dir / 'NS3_Subtype_RAS_Profiles.xlsx',
        'NS5A': args.ns5a_output_dir / 'NS5A_Subtype_RAS_Profiles.xlsx',
        'NS5B': args.ns5b_output_dir / 'NS5B_Subtype_RAS_Profiles.xlsx',
    }
    gene_tokens_by_runner_gene = {'NS3': ('NS3', 'NS3'), 'NS5A': ('NS5A_NTD', 'NS5A'), 'NS5B': ('NS5B', 'NS5B')}
    gene_tokens = [gene_tokens_by_runner_gene[args.gene]]
    comparison_gene, profile_gene = gene_tokens[0]
    if args.comparison_workbook:
        comparison_paths[comparison_gene] = args.comparison_workbook
    if args.ras_workbook:
        profile_paths[profile_gene] = args.ras_workbook
    allowed_subtypes = None
    if args.combined_profile_workbook:
        allowed_subtypes = combined_profile_subtypes(args.combined_profile_workbook)
    report_genes = [gene_tokens[0][0]]
    required = [comparison_paths[report_genes[0]], profile_paths[gene_tokens[0][1]]]
    missing = [str(path) for path in required if not path.is_file()]
    if missing:
        print('shared_report_status=waiting_for_gene_inputs')
        print('missing_inputs=' + ';'.join(missing))
        return
    validate_reported_mutations(comparison_paths, profile_paths, gene_tokens)
    readme: dict[str,dict[str,set[str]]] = {}
    for gene,token in gene_tokens:
        path = comparison_paths[gene]
        values = profiles(profile_paths[token])
        wb=load_workbook(path); ws=wb.active
        col=next((cell.column for cell in ws[2] if cell.value=='Mutations'), ws.max_column+1)
        for row in range(1,ws.max_row+1):
            if ws.cell(row,1).value=='Gene': ws.cell(row,col).value='Mutations'; ws.cell(row,col).font=Font(bold=True)
            if ws.cell(row,4).value!='Consensus': continue
            refrow=row-1; gt=str(ws.cell(row,2).value or '').removeprefix('GT'); subtype=str(ws.cell(row,3).value or '').lower(); parts=[]; plain=[]
            if allowed_subtypes is not None and f'GT{gt}_{subtype}' not in allowed_subtypes:
                continue
            for c in range(5,col):
                cons=ws.cell(row,c).value
                if not cons: continue
                ref=ws.cell(refrow,c).value; pos=ws.cell(2,c).value
                if not ref or not pos: continue
                mutation=f'{ref}{pos}{cons}'; pct=values[(gt,subtype,int(pos),str(cons))]
                parts.extend([mutation, TextBlock(InlineFont(vertAlign='subscript'), pct), ', '])
                plain.append(f'{mutation} {pct}%')
            if parts:
                parts.pop(); ws.cell(row,col).value=CellRichText(*parts)
                readme.setdefault(gene,{}).setdefault(f'GT{gt}_{subtype}',set()).update(plain)
        ws.column_dimensions[ws.cell(1,col).column_letter].width=50; wb.save(path)
    write_summary_documents(output_root, readme, report_genes)
    publish_shared_report(comparison_paths, output_root, args.shared_report_dir, report_genes)

if __name__=='__main__': main()
